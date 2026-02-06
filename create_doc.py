"""
Generate Word Document for Nova Project Documentation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
title = doc.add_heading('Nova - Amazon Voice Shopping Assistant', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Project Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1: Project Understanding
doc.add_heading('1. Project Understanding', level=1)
doc.add_paragraph(
    'This project is a voice-controlled Amazon shopping assistant named "Nova". '
    'It enables users to shop on Amazon.in using natural voice commands through a web-based interface. '
    'The system combines real-time voice AI with browser automation to create a hands-free shopping experience.'
)
doc.add_paragraph(
    '💡 Key Insight: Nova acts as a voice-to-browser bridge: users speak commands, '
    'the AI interprets them, and automated browser actions perform the shopping tasks on Amazon.'
)

# Section 2: What We're Building
doc.add_heading('2. What We Are Building', level=1)
doc.add_heading('The Product: Voice-Powered Amazon Shopping Assistant', level=2)

doc.add_paragraph('The system consists of these key components:')

# Components table
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'File'
hdr_cells[2].text = 'Purpose'

data = [
    ('Voice Agent', 'agent.py', 'Handles voice I/O via LiveKit + Gemini Realtime API'),
    ('Web Server', 'server.py', 'Flask server for client tokens and UI hosting'),
    ('Browser Controller', 'browser_controller.py', 'Playwright-based Amazon automation'),
    ('Shopping Tools', 'automation_tools.py', 'Voice-triggered shopping functions'),
    ('Configuration', 'config.py', 'System settings and AI prompts'),
    ('Web UI', 'templates/index.html', 'Modern voice interface with visualizations'),
]

for i, (comp, file, purpose) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = comp
    row[1].text = file
    row[2].text = purpose

doc.add_paragraph()

# Section 3: Development Output
doc.add_heading('3. Development Output', level=1)

doc.add_heading('3.1 A Web Interface', level=2)
doc.add_paragraph('• Modern, animated UI with connection status')
doc.add_paragraph('• Real-time audio visualizer')
doc.add_paragraph('• Transcript display of conversations')
doc.add_paragraph('• Quick-action suggestion chips')

doc.add_heading('3.2 Voice Shopping Capabilities', level=2)
doc.add_paragraph('Users can perform these actions via voice:')

# Voice commands table
cmd_table = doc.add_table(rows=11, cols=2)
cmd_table.style = 'Table Grid'
cmd_hdr = cmd_table.rows[0].cells
cmd_hdr[0].text = 'Voice Command'
cmd_hdr[1].text = 'Action'

commands = [
    ('"Open Amazon"', 'Opens Amazon.in in browser'),
    ('"Search for wireless earbuds"', 'Searches and displays results'),
    ('"Show me the first one"', 'Opens product details'),
    ('"Add it to my cart"', 'Adds item to shopping cart'),
    ('"What\'s in my cart?"', 'Opens cart view'),
    ('"Proceed to checkout"', 'Starts checkout process'),
    ('"Filter by Prime"', 'Applies Prime delivery filter'),
    ('"Read reviews"', 'Reads customer reviews aloud'),
    ('"Scroll down"', 'Shows more products'),
    ('"Go back"', 'Returns to previous page'),
]

for i, (cmd, action) in enumerate(commands, 1):
    row = cmd_table.rows[i].cells
    row[0].text = cmd
    row[1].text = action

doc.add_paragraph()

doc.add_heading('3.3 An Automated Browser', level=2)
doc.add_paragraph('A visible Chromium browser window that:')
doc.add_paragraph('• Performs searches automatically')
doc.add_paragraph('• Navigates product pages')
doc.add_paragraph('• Handles cart operations')
doc.add_paragraph('• Responds in real-time to voice commands')

# Section 4: What Nova Does
doc.add_heading('4. What Nova Does', level=1)

doc.add_heading('Functional Flow', level=2)
doc.add_paragraph(
    '1. User opens web interface and clicks "Start Shopping"\n'
    '2. WebRTC connection established via LiveKit\n'
    '3. Nova greets the user with a welcome message\n'
    '4. User speaks a shopping command (e.g., "Search for laptops")\n'
    '5. Gemini AI processes the intent and identifies the action\n'
    '6. Appropriate tool function is called (e.g., search_product)\n'
    '7. Playwright controls Chromium browser to execute action on Amazon\n'
    '8. Results are captured and Nova speaks the response\n'
    '9. Process repeats for next command'
)

doc.add_heading('Core Functionality', level=2)
doc.add_paragraph('1. Voice Recognition: Captures user speech through browser microphone')
doc.add_paragraph('2. Intent Understanding: Gemini AI interprets shopping requests')
doc.add_paragraph('3. Tool Execution: Calls appropriate automation functions')
doc.add_paragraph('4. Browser Control: Playwright controls Chromium to navigate Amazon')
doc.add_paragraph('5. Voice Response: Nova speaks results and asks follow-up questions')

# Section 5: Technology Stack
doc.add_heading('5. Technology Stack', level=1)

doc.add_heading('Backend Technologies', level=2)
tech_table = doc.add_table(rows=7, cols=3)
tech_table.style = 'Table Grid'
tech_hdr = tech_table.rows[0].cells
tech_hdr[0].text = 'Technology'
tech_hdr[1].text = 'Version'
tech_hdr[2].text = 'Purpose'

techs = [
    ('Python', '3.10+', 'Core programming language'),
    ('Flask', '3.0+', 'Web server framework'),
    ('LiveKit', '0.18+', 'Real-time voice/audio infrastructure'),
    ('LiveKit Agents', '0.12+', 'Agent framework for voice AI'),
    ('Gemini Realtime API', 'Latest', 'Voice AI + natural language understanding'),
    ('Playwright', '1.40+', 'Browser automation'),
]

for i, (tech, ver, purpose) in enumerate(techs, 1):
    row = tech_table.rows[i].cells
    row[0].text = tech
    row[1].text = ver
    row[2].text = purpose

doc.add_paragraph()

doc.add_heading('Frontend Technologies', level=2)
front_table = doc.add_table(rows=5, cols=2)
front_table.style = 'Table Grid'
front_hdr = front_table.rows[0].cells
front_hdr[0].text = 'Technology'
front_hdr[1].text = 'Purpose'

fronts = [
    ('HTML5', 'Web interface structure'),
    ('CSS3', 'Modern styling with animations'),
    ('JavaScript', 'LiveKit client SDK integration'),
    ('LiveKit Client SDK', 'Browser-side voice handling'),
]

for i, (tech, purpose) in enumerate(fronts, 1):
    row = front_table.rows[i].cells
    row[0].text = tech
    row[1].text = purpose

doc.add_paragraph()

doc.add_heading('External Services', level=2)
ext_table = doc.add_table(rows=4, cols=2)
ext_table.style = 'Table Grid'
ext_hdr = ext_table.rows[0].cells
ext_hdr[0].text = 'Service'
ext_hdr[1].text = 'Purpose'

exts = [
    ('LiveKit Cloud', 'WebRTC voice streaming infrastructure'),
    ('Google Gemini', 'Real-time voice AI model'),
    ('Amazon.in', 'Target shopping platform'),
]

for i, (svc, purpose) in enumerate(exts, 1):
    row = ext_table.rows[i].cells
    row[0].text = svc
    row[1].text = purpose

# Section 6: Application Flow
doc.add_heading('6. Application Flow', level=1)

doc.add_heading('System Architecture', level=2)
doc.add_paragraph(
    'The application consists of four main layers:\n\n'
    '1. CLIENT (Web Browser)\n'
    '   • index.html + app.js - User interface\n'
    '   • Microphone - Captures voice input\n'
    '   • Speaker - Plays voice responses\n\n'
    '2. SERVER (Flask - Port 5000)\n'
    '   • server.py - Handles token generation\n'
    '   • templates/ - HTML templates\n'
    '   • static/ - CSS and JavaScript files\n\n'
    '3. AGENT (Nova Voice Agent)\n'
    '   • agent.py - Main agent logic\n'
    '   • automation_tools.py - Shopping functions\n'
    '   • browser_controller.py - Browser automation\n\n'
    '4. EXTERNAL SERVICES\n'
    '   • LiveKit Cloud - Voice streaming\n'
    '   • Google Gemini - AI processing\n'
    '   • Amazon.in - Shopping platform'
)

doc.add_heading('Startup Flow', level=2)
doc.add_paragraph(
    '1. Start Flask Server (start_server.bat)\n'
    '   • Runs on localhost:5000\n'
    '   • Serves web UI and generates LiveKit tokens\n\n'
    '2. Start Voice Agent (start_agent.bat)\n'
    '   • Connects to LiveKit room\n'
    '   • Waits for users to join\n\n'
    '3. User Opens Web UI\n'
    '   • Navigates to http://localhost:5000\n'
    '   • Clicks "Start Shopping" button\n\n'
    '4. Voice Session Begins\n'
    '   • WebRTC connection established via LiveKit\n'
    '   • Nova greets the user\n'
    '   • Chromium browser launches for Amazon automation'
)

# Section 7: File Structure
doc.add_heading('7. File Structure', level=1)

doc.add_paragraph(
    'Voice_Assistant/\n'
    '├── agent.py              # Main voice agent (LiveKit + Gemini)\n'
    '├── server.py             # Flask web server\n'
    '├── config.py             # Configuration & AI prompts\n'
    '├── automation_tools.py   # Voice-triggered shopping tools\n'
    '├── browser_controller.py # Playwright browser automation\n'
    '├── requirements.txt      # Python dependencies\n'
    '├── .env                  # API keys (LiveKit, Gemini)\n'
    '├── setup.bat             # Environment setup script\n'
    '├── start_server.bat      # Starts Flask server\n'
    '├── start_agent.bat       # Starts voice agent\n'
    '├── templates/\n'
    '│   └── index.html        # Web UI template\n'
    '├── static/\n'
    '│   ├── app.js            # Client-side JavaScript\n'
    '│   └── styles.css        # UI styling\n'
    '└── venv/                 # Python virtual environment'
)

# Section 8: Summary
doc.add_heading('8. Summary', level=1)

summary_table = doc.add_table(rows=8, cols=2)
summary_table.style = 'Table Grid'

summaries = [
    ('What', 'Voice-controlled Amazon shopping assistant'),
    ('Name', 'Nova'),
    ('How', 'Combines Gemini Realtime API for voice AI with Playwright for browser automation'),
    ('Interface', 'Modern web-based UI with audio visualizations'),
    ('Target', 'Amazon.in e-commerce platform'),
    ('Capabilities', 'Search, browse, add to cart, filter, read reviews, checkout'),
    ('Experience', 'Hands-free shopping through natural conversation'),
]

sum_hdr = summary_table.rows[0].cells
sum_hdr[0].text = 'Aspect'
sum_hdr[1].text = 'Description'

for i, (aspect, desc) in enumerate(summaries, 1):
    row = summary_table.rows[i].cells
    row[0].text = aspect
    row[1].text = desc

doc.add_paragraph()

# How to Run
doc.add_heading('How to Run the Application', level=2)
doc.add_paragraph(
    '1. Run start_server.bat to start the web server\n'
    '2. Run start_agent.bat to start the voice agent\n'
    '3. Open http://localhost:5000 in your browser\n'
    '4. Click "Start Shopping" and begin speaking!'
)

# Save the document
doc.save('d:/Voice_Assistant/Nova_Project_Documentation.docx')
print("✅ Word document created successfully!")
print("📄 Saved to: d:/Voice_Assistant/Nova_Project_Documentation.docx")
